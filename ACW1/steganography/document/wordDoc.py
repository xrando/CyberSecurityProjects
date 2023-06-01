from docx import Document
from docx.shared import RGBColor
import numpy as np

# converting the data to binary format as string
# checks the type of the data and perform the conversion accordingly
def toBinary(data):
    if isinstance(data, str):
        return ''.join([format(ord(i), "08b") for i in data])
    elif isinstance(data, bytes):
        return ''.join([format(i, "08b") for i in data])
    elif isinstance(data, np.ndarray) or isinstance(data, list) or isinstance(data, tuple):
        return [format(i, "08b") for i in data]
    elif isinstance(data, int) or isinstance(data, np.uint8):
        return format(data, "08b")
    else:
        raise TypeError("Type not supported.")

class fontcolourSteganography:
    # encoding function
    def encode(self, filePath, payload_file, bit):
        # opens the Word document 
        document = Document(filePath)
        
        # open & read the content of the payload file
        with open(payload_file, 'r', encoding='utf-8') as file:
            payload = file.read()
            
        # adds "=====" to the payload to mark the end of the message
        payload += "====="
        # converts payload to binary using the binary func above
        binary_msg = toBinary(payload)
        # calculates the length of the binary message
        length = len(binary_msg)
        # check max length
        max_length = sum(len(paragraph.text) for paragraph in document.paragraphs)
        # if the length exceeds then 'ValueError' will be raised
        if length > max_length:
            raise ValueError(f"Message length ({length}) exceeded {max_length}, please adjust the message or bit.")

        index = 0
        # iterates over each paragraph in the document
        for paragraph in document.paragraphs:
            # if the 'index' pos exceeds the length of message then breaks the loop
            if index >= length:
                break
            # retrieve the runs of segments of text with the same formatting within the paragraph
            runs = paragraph.runs
            # clears the paragraph to remove existing text
            paragraph.clear()
            # for each run iterate over each character
            for run in runs:
                text = run.text
                color = toBinary(run.font.color.rgb) if run.font.color.rgb is not None else ['00000000', '00000000', '00000000']
                for _, char in enumerate(text):
                    # if index eceeds the length of msg, adds char to the paragraph and continues to next char
                    if index >= length:
                        paragraph.add_run(char)
                        continue
                    # if not add char to the run
                    run = paragraph.add_run(char)
                    # retrieve the appropriate bits from binary msg or the color values and modify them
                    for i in range(3):
                        data = "".join(binary_msg[index:index+bit]) if index < length else color[i][-bit:]
                        index += bit
                        color[i] = color[i][:-bit] + data
                    # sets the modified color using'RGBColor'
                    run.font.color.rgb = RGBColor(int(color[0], 2), int(color[1], 2), int(color[2], 2))
        # returns the modified document at the end
        return document

    # decoding function
    def decode(self, filePath, bit, outputFilePath):
        # opens the Word document 
        document = Document(filePath)
        # initialize the variables
        binary_msg = ""
        payload = ""
        # iterates over each paragraph in the document
        for paragraph in document.paragraphs:
            # for each run in the document
            for run in paragraph.runs:
                # if the color of the run is None
                if run.font.color.rgb is None:
                    continue
                # converts each color component to binary and appends the lsb to the binary msg
                for color_component in run.font.color.rgb:
                    binary_color = toBinary(color_component)
                    binary_msg += binary_color[-bit:]
                    # if the length reaches 8 or more, converts the first 8 bits to a char and append to payload
                    if len(binary_msg) >= 8:
                        payload += chr(int(binary_msg[:8], 2))
                        # if payload ends with "======", breaks the loop as it is the end of the message
                        if payload.endswith("====="):
                            break
                        # updates the binary message by removing first 8 bits
                        binary_msg = binary_msg[8:]

        # after processing all runs => writes payload to specified output file
        with open(outputFilePath, 'w', encoding='utf-8') as file:
            file.write(payload[:-5])

        # returns payload (which will exclude the marker "=====" which signifies the end of the message)
        return payload[:-5]

# creating an instance
obj = fontcolourSteganography()

# specify encoding file path
encodeFilePath = 'cover.docx'

# specifying the payload content
payload_file = 'payload.txt'

# specifying the no.of bits
bit = 5

# specifying the output file path
outputFilePath = 'encoded_payload.txt'

# specifying the save file path
saveFilePath = "C:/Users/Jovian Low/Desktop/SIT UoG/SIT Y1Tri3/Cybersecurity/Jovian part/encoded.docx"

# using the encode method and save the modified document
obj.encode(encodeFilePath, payload_file, bit).save(saveFilePath)

# prints the decoded payload out
print(obj.decode(saveFilePath, bit, outputFilePath))
