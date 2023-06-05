from docx import Document

class HiddenTextSteganography:
    # Encoding function
    def encode(self, file_path, payload_file, save_file_path):
        # Open the Word document
        document = Document(file_path)

        # Open and read the content of the payload file
        with open(payload_file, 'r', encoding='utf-8') as file:
            payload = file.read()

        # Add hidden text to the document
        document.add_paragraph().add_run(payload).font.hidden = True

        # Save the modified document
        document.save(save_file_path)

        return save_file_path

    # Decoding function
    def decode(self, file_path, output_file_path):
        # Open the Word document
        document = Document(file_path)

        # Initialize the payload variable
        payload = ""

        # Iterate over each paragraph in the document
        for paragraph in document.paragraphs:
            # Check if the run has hidden text
            for run in paragraph.runs:
                if run.font.hidden:
                    # Append the text to the payload
                    payload += run.text

        # Write the payload to the output file
        with open(output_file_path, 'w', encoding='utf-8') as file:
            file.write(payload)

        return payload


if __name__ == "__main__":
    # Create an instance
    obj = HiddenTextSteganography()

    # Specify the encoding file path
    encode_file_path = 'files/cover.docx'

    # Specify the payload content
    payload_file = 'files/longpayload.txt'

    # Specify the save file path
    save_file_path = 'files/encoded_payload.docx'

    # Using the encode method and save the modified document
    encoded_file_path = obj.encode(encode_file_path, payload_file, save_file_path)

    # Print the encoded file path
    print("Encoded file:", encoded_file_path)

    # Specify the output file path
    output_file_path = 'files/decoded_payload.txt'

    # Using the decode method and save the decoded payload to the output file
    decoded_payload = obj.decode(encoded_file_path, output_file_path)

    # Print the decoded payload
    print("Decoded payload:", decoded_payload)